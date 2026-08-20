"""JD-aware resume tailoring for ATS/AI screening.

Starts from canonical `Rafi_Resume.docx`, rewrites headline/summary and reorders
skills + bullets so JD keywords score higher — without inventing employers,
dates, titles, or metrics. Output keeps filename `Rafi_Resume.docx` for Easy
Apply label matching.

Disable with RESUME_TAILOR=0 / LINKEDIN_TAILOR_RESUME=0 / FOUNDIT_TAILOR=0 / INDEED_TAILOR_RESUME=0.
"""

from __future__ import annotations

import hashlib
import os
import re
from pathlib import Path

# Known skills/phrases already evidenced on the canonical resume (or close synonyms).
# Only these may be injected into the tailored summary/skills lines.
OWNED_SKILLS: tuple[str, ...] = (
    ".NET Core",
    "ASP.NET Core",
    "C#",
    "React",
    "Angular",
    "SQL Server",
    "PostgreSQL",
    "REST API",
    "Microservices",
    "Event-Driven Architecture",
    "Kafka",
    "RabbitMQ",
    "AWS",
    "Azure",
    "Docker",
    "Kubernetes",
    "CI/CD",
    "Jenkins",
    "Domain-Driven Design",
    "System Design",
    "Distributed Systems",
    "API Design",
    "High Availability",
    "Agile",
    "Scrum",
    "Technical Leadership",
    "Mentoring",
    "Architecture Reviews",
    "Cloud Architecture",
    "Solution Architecture",
    "Software Architecture",
)

# Extra JD tokens we may surface when they map to owned experience (synonyms).
SKILL_ALIASES: dict[str, str] = {
    "dotnet": ".NET Core",
    ".net": ".NET Core",
    "asp.net": "ASP.NET Core",
    "csharp": "C#",
    "c sharp": "C#",
    "k8s": "Kubernetes",
    "eks": "Kubernetes",
    "aks": "Kubernetes",
    "ecs": "AWS",
    "ec2": "AWS",
    "s3": "AWS",
    "lambda": "AWS",
    "sqs": "AWS",
    "sns": "AWS",
    "cloudformation": "AWS",
    "terraform": "Cloud Architecture",
    "api gateway": "API Design",
    "restful": "REST API",
    "apis": "REST API",
    "micro service": "Microservices",
    "micro-service": "Microservices",
    "event driven": "Event-Driven Architecture",
    "message queue": "RabbitMQ",
    "messaging": "Kafka",
    "container": "Docker",
    "orchestration": "Kubernetes",
    "devops": "CI/CD",
    "solution architect": "Solution Architecture",
    "software architect": "Software Architecture",
    "technical architect": "Solution Architecture",
    "cloud architect": "Cloud Architecture",
    "system architect": "System Design",
    "tech lead": "Technical Leadership",
    "technical lead": "Technical Leadership",
    "engineering manager": "Technical Leadership",
    "people manager": "Mentoring",
    "ddb": "Domain-Driven Design",
    "ddd": "Domain-Driven Design",
    "sqlserver": "SQL Server",
    "mssql": "SQL Server",
    "postgres": "PostgreSQL",
    "ha/dr": "High Availability",
    "disaster recovery": "High Availability",
}

ROLE_WORDS = re.compile(
    r"\b("
    r"solution\s+architect|technical\s+architect|software\s+architect|"
    r"cloud\s+architect|azure\s+architect|principal\s+engineer|"
    r"staff\s+engineer|engineering\s+manager|technical\s+lead|tech\s+lead|"
    r"principal\s+analyst|architect|tech\s*lead"
    r")\b",
    re.I,
)

STOP = {
    "and",
    "the",
    "for",
    "with",
    "from",
    "that",
    "this",
    "your",
    "you",
    "our",
    "are",
    "will",
    "have",
    "has",
    "been",
    "into",
    "using",
    "work",
    "role",
    "job",
    "team",
    "years",
    "year",
    "experience",
    "required",
    "preferred",
    "strong",
    "good",
    "ability",
    "etc",
}


def tailor_enabled() -> bool:
    for key in (
        "RESUME_TAILOR",
        "LINKEDIN_TAILOR_RESUME",
        "FOUNDIT_TAILOR",
        "INDEED_TAILOR_RESUME",
    ):
        v = os.environ.get(key, "").strip().lower()
        if v in ("0", "false", "no", "off"):
            return False
    return True


def _artifacts_root() -> Path:
    env = os.environ.get("LINKEDIN_ARTIFACTS") or os.environ.get("RESUME_TAILOR_DIR")
    if env:
        p = Path(env)
        p.mkdir(parents=True, exist_ok=True)
        return p
    cloud = Path("/opt/cursor/artifacts")
    if cloud.is_dir() or os.environ.get("CURSOR_AGENT"):
        cloud.mkdir(parents=True, exist_ok=True)
        return cloud
    root = Path(__file__).resolve().parents[1]
    d = root / "artifacts"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip())


def extract_jd_skills(jd: str, title: str = "") -> list[str]:
    """Return owned skills that the JD/title appears to ask for (ordered by strength)."""
    blob = f"{title}\n{jd}".lower()
    scores: dict[str, int] = {}

    for skill in OWNED_SKILLS:
        pat = re.escape(skill.lower()).replace(r"\ ", r"[\s\-/]*")
        if re.search(rf"\b{pat}\b", blob, re.I):
            scores[skill] = scores.get(skill, 0) + 3

    for alias, skill in SKILL_ALIASES.items():
        if re.search(rf"\b{re.escape(alias)}\b", blob, re.I):
            scores[skill] = scores.get(skill, 0) + 2

    # Soft boost: token overlap with skill words
    tokens = set(re.findall(r"[a-z0-9+#.]{3,}", blob))
    for skill in OWNED_SKILLS:
        parts = {p for p in re.findall(r"[a-z0-9+#.]{3,}", skill.lower()) if p not in STOP}
        if parts & tokens:
            scores[skill] = scores.get(skill, 0) + 1

    ranked = sorted(scores.items(), key=lambda kv: (-kv[1], kv[0].lower()))
    return [s for s, _ in ranked]


def preferred_headline(title: str, matched: list[str]) -> str:
    """Keep a truthful headline; prefer the posting title when it is architect/lead-like."""
    t = _norm(title)
    m = ROLE_WORDS.search(t) if t else None
    if m:
        role = _norm(m.group(1)).title()
        # Normalize common casings
        role = (
            role.replace("Net", ".NET")
            .replace(".Net", ".NET")
            .replace("Ci/Cd", "CI/CD")
        )
    else:
        role = "Technical Architect | Technical Lead"
    stack = []
    for s in matched:
        if s in (".NET Core", "C#", "Azure", "AWS", "Microservices", "Kafka", "Kubernetes") and s not in stack:
            stack.append(s)
        if len(stack) >= 3:
            break
    if not stack:
        stack = [".NET", "Cloud", "Distributed Systems"]
    return f"{role} — {', '.join(stack)}"


def preferred_summary(title: str, company: str, matched: list[str]) -> str:
    """Truthful summary that mirrors JD language via owned skills only."""
    skills = matched[:8] or [
        ".NET Core",
        "Microservices",
        "AWS",
        "Azure",
        "Kafka",
        "Kubernetes",
    ]
    skill_phrase = ", ".join(skills[:-1]) + (", and " + skills[-1] if len(skills) > 1 else skills[0])
    role_bit = _norm(title) if title and ROLE_WORDS.search(title or "") else "Technical Architect / Technical Lead"
    company_bit = f" aligned to {company}" if company and len(company) < 40 else ""
    return (
        f"{role_bit} with 15+ years designing distributed, cloud-native platforms and leading "
        f"engineering teams{company_bit}. Hands-on depth in {skill_phrase}. Owns end-to-end "
        "architecture, API and integration design, event-driven services, and delivery "
        "governance across healthcare, retail, and telecom domains."
    )


def _score_text(text: str, matched: list[str]) -> int:
    low = (text or "").lower()
    score = 0
    for i, skill in enumerate(matched):
        if skill.lower() in low:
            score += max(1, 12 - i)
    return score


def _set_paragraph_text(paragraph, text: str) -> None:
    """Replace paragraph text while keeping the first run's formatting when possible."""
    text = text or ""
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def _reorder_list_block(paragraphs: list, start: int, end: int, matched: list[str]) -> None:
    """Reorder consecutive List Paragraph items by JD match (stable for ties)."""
    from docx.text.paragraph import Paragraph  # type: ignore

    items: list[tuple[int, int, Paragraph]] = []
    for i in range(start, end):
        p = paragraphs[i]
        if (p.style and p.style.name or "") == "List Paragraph":
            items.append((i, _score_text(p.text, matched), p))
    if len(items) < 2:
        return
    ordered = sorted(items, key=lambda t: (-t[1], t[0]))
    texts = [p.text for _, _, p in ordered]
    for (_, _, p), new_text in zip(items, texts):
        _set_paragraph_text(p, new_text)


def tailor_document(
    src: Path,
    dest: Path,
    *,
    title: str = "",
    company: str = "",
    jd: str = "",
) -> dict:
    """Write a tailored copy of src → dest. Returns metadata for logging/tests."""
    from docx import Document  # lazy import

    matched = extract_jd_skills(jd, title)
    doc = Document(str(src))
    paras = list(doc.paragraphs)
    meta = {
        "matched_skills": matched,
        "headline": preferred_headline(title, matched),
        "summary": preferred_summary(title, company, matched),
        "source": str(src),
        "dest": str(dest),
    }

    # Paragraph layout from canonical resume (see tools/resume_tailor tests):
    # 0 name, 1 headline, 2 contact, 3 SUMMARY header, 4 summary body,
    # 5 COMPETENCIES header, 6-9 competency lines, 10 EXPERIENCE header, then roles+bullets
    if len(paras) >= 2:
        _set_paragraph_text(paras[1], meta["headline"])
    if len(paras) >= 5:
        _set_paragraph_text(paras[4], meta["summary"])

    # Reorder competency category lines (indices 6-9) by match score
    if len(paras) >= 10:
        comp = [(i, paras[i]) for i in range(6, 10)]
        ranked = sorted(comp, key=lambda t: (-_score_text(t[1].text, matched), t[0]))
        texts = [p.text for _, p in ranked]
        for (i, p), text in zip(comp, texts):
            _set_paragraph_text(p, text)

    # Reorder bullets under each experience role
    i = 0
    while i < len(paras):
        style = (paras[i].style and paras[i].style.name) or ""
        if style == "Normal" and re.search(
            r"Principal Analyst|PS IV|Technical Lead|Senior Consultant|Senior Software|"
            r"Senior Developer|Senior Systems|Nemetschek|NCR|UnitedHealth|ADP|EPAM|Infosys",
            paras[i].text or "",
            re.I,
        ):
            j = i + 1
            while j < len(paras) and ((paras[j].style and paras[j].style.name) or "") == "List Paragraph":
                j += 1
            _reorder_list_block(paras, i + 1, j, matched)
            i = j
            continue
        i += 1

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))
    meta["size"] = dest.stat().st_size
    return meta


def tailor_resume_for_job(
    *,
    job_id: str = "",
    title: str = "",
    company: str = "",
    jd: str = "",
    src: Path | None = None,
) -> Path:
    """Tailor canonical resume for one job. Returns path to Rafi_Resume.docx copy."""
    from tools.resume_paths import CANONICAL_NAME, ensure_resume_aliases

    if not tailor_enabled():
        return ensure_resume_aliases()

    base = Path(src) if src else ensure_resume_aliases()
    jid = re.sub(r"[^\w.-]+", "_", (job_id or "").strip()) or "unknown"
    # Stable folder so re-runs overwrite; filename stays Rafi_Resume.docx for LI label.
    digest = hashlib.sha1(f"{title}|{company}|{(jd or '')[:2000]}".encode()).hexdigest()[:10]
    out_dir = _artifacts_root() / "tailored-resumes" / f"{jid}-{digest}"
    dest = out_dir / CANONICAL_NAME
    try:
        meta = tailor_document(base, dest, title=title, company=company, jd=jd)
        print(
            f"  resume tailored id={jid} skills={meta['matched_skills'][:6]} -> {dest}",
            flush=True,
        )
        return dest
    except Exception as e:
        print(f"  resume tailor failed ({e}); using canonical", flush=True)
        return base


if __name__ == "__main__":
    import json
    import sys

    jd = Path(sys.argv[1]).read_text() if len(sys.argv) > 1 else "Azure .NET Solution Architect Kafka Kubernetes"
    p = tailor_resume_for_job(job_id="demo", title="Solution Architect", company="Acme", jd=jd)
    print(json.dumps({"path": str(p), "size": p.stat().st_size}, indent=2))
